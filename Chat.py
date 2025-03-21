import streamlit as st
from io import BytesIO
from streamlit_mic_recorder import mic_recorder
import re
from docx import Document
import warnings
import io
import time
import codecs
from Logger import log
from Utils import get_uuid, elapsed, autoplay_audio, local_css, get_prompt
from Session import (
    get_session,
    update_active_users,
    get_active_user_count,
    log_session,
    push_session_log,
)
from OpenAIClient import speech_to_text, text_to_speech
from Settings import settings

if settings["parameters"]["model"].startswith("claude"):
    from AnthropicClient import get_response
else:
    from OpenAIClient import get_response

warnings.filterwarnings("ignore", category=DeprecationWarning)


def show_download(container):
    document = create_transcript_document()
    container.download_button(
        label="Download Transcript",
        icon=":material/download:",
        data=document,
        file_name="Transcript.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def create_transcript_document():
    doc = Document()
    doc.add_heading("Conversation Transcript\n", level=1)

    for message in st.session_state.messages[1:]:
        if message["role"] == "user":
            p = doc.add_paragraph()
            p.add_run(settings["user_name"] + ": ").bold = True
            p.add_run(message["content"])
        else:
            p = doc.add_paragraph()
            p.add_run(settings["assistant_name"] + ": ").bold = True
            p.add_run(message["content"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Session Initialization
def init_session():
    st.session_state["start_time"] = time.time()
    if "messages" not in st.session_state:
        st.session_state["messages"] = [{"role": "system", "content": get_prompt()}]
    # Cache large prompt to cut down first response time.
    get_response(st.session_state.messages)
    st.session_state["manual_input"] = None
    st.session_state["text_chat_enabled"] = False
    st.session_state["end_session_button_clicked"] = False
    st.session_state["download_transcript"] = False
    autoplay_audio(open("assets/unlock.mp3", "rb").read())
    update_active_users()
    log.info(
        f"Session Start: {time.time()-st.session_state.start_time:.2f} seconds, {get_session()}"
    )


def setup_sidebar():
    st.sidebar.header("Chat with " + settings["assistant_name"])
    # Show Profile
    with st.sidebar.container(border=True):
        for key, val in settings["sidebar"].items():
            if re.search(r"(jpg|png|webp)$", val):
                st.image(val)
            else:
                st.subheader(f"{key.replace("_", " ")}: {val}")

    def toggle_text_chat():
        st.session_state.text_chat_enabled = not st.session_state.text_chat_enabled

    container = st.sidebar.container(border=True)
    con1 = container.container()
    con2 = container.container(border=True)
    con3 = container.container()
    con4 = container.empty()
    con2.toggle(
        ":material/keyboard: Enable Text Chat",
        value=st.session_state.text_chat_enabled,
        on_change=toggle_text_chat,
    )
    return con1, con3, con4


def show_messages():
    for message in st.session_state.messages[1:]:
        name = (
            settings["user_name"]
            if message["role"] == "user"
            else settings["assistant_name"]
        )
        avatar = (
            settings["user_avatar"]
            if message["role"] == "user"
            else settings["assistant_avatar"]
        )
        with st.chat_message(name, avatar=avatar):
            st.markdown(message["content"])


def handle_audio_input(container):
    with container:
        if audio := mic_recorder(
            start_prompt="🎙 Record",
            stop_prompt="📤 Stop",
            just_once=True,
            use_container_width=True,
            format="wav",
            key="recorder",
        ):
            return speech_to_text(audio)


def process_user_query(user_query, container):
    # Store the user's query into the history
    st.session_state.messages.append({"role": "user", "content": user_query.strip()})
    # Display the user's query
    if st.session_state.text_chat_enabled:
        with st.chat_message(
            settings["user_name"],
            avatar=settings["user_avatar"],
        ):
            st.markdown(user_query)
    if st.session_state.manual_input:
        response = get_response(
            st.session_state.messages,
            settings["parameters"]["feedback_model"],
            settings["parameters"]["feedback_temperature"],
        )
    elif st.session_state.end_session_button_clicked:
        response = get_response(
            st.session_state.messages,
            temperature=settings["parameters"]["feedback_temperature"],
        )
    else:
        response = get_response(st.session_state.messages)
    response = response.strip()
    st.session_state.messages.append({"role": "assistant", "content": response})
    if not st.session_state.end_session_button_clicked:
        if audio := text_to_speech(response, instructions="Speak like a nervous and guarded teenager."):
            autoplay_audio(audio, container)

    if st.session_state.text_chat_enabled:
        with st.chat_message(
            settings["assistant_name"],
            avatar=settings["assistant_avatar"],
        ):
            st.markdown(response)

    log_session()


st.set_page_config(
    page_title="Chat | " + settings["title"],
    page_icon=":material/chat:",
    layout="wide",
    initial_sidebar_state="expanded",
)
st.title("Chat | " + settings["title"])
if "role" not in st.session_state:
    st.switch_page("Login.py")

# Inject CSS for custom styles
local_css("style.css")
if "text_chat_enabled" not in st.session_state:
    init_session()
container1, container3, container4 = setup_sidebar()
if st.session_state.text_chat_enabled:
    show_messages()
else:
    st.markdown(
        """
        You are in voice chat-only mode, which disables text input and hides the conversation history.

        When you are ready, click **🎙 Record**, allow microphone access if prompted, speak when the button changes to **📤 Stop**, then click **📤 Stop** when you are done speaking.

        If you experience issues with voice chat, click **Enable Text Chat** in the left panel.
    """
    )

# Check if there's a manual input and process it
if st.session_state.manual_input:
    container3.button(
        "🤔 Generating Feedback...", icon=":material/feedback:", disabled=True
    )
    user_query = st.session_state.manual_input
else:
    if st.session_state.end_session_button_clicked:
        user_query = st.chat_input(
            "Ask questions about your feedback below or click 'Start Over' in the left panel."
        )
    else:
        user_query = st.chat_input(
            "Click 'End Session' Button in the Left Panel to Receive Feedback and Download Transcript.",
            disabled=not st.session_state.text_chat_enabled,
        )
        if transcript := handle_audio_input(container1):
            user_query = transcript

if user_query:
    process_user_query(user_query, container4)
    if st.session_state.manual_input:
        st.session_state.manual_input = None
        push_session_log()
        st.rerun()

# Handle end session
if not st.session_state.end_session_button_clicked:
    if len(st.session_state.messages) > 1:
        if container3.button(
            "End Session",
            icon=":material/call_end:",
            disabled=st.session_state.end_session_button_clicked,
        ):
            st.session_state.end_session_button_clicked = True
            st.session_state.text_chat_enabled = True
            log.info(
                f"Session end: {elapsed(st.session_state.start_time)} {get_session()}"
            )
            st.session_state.download_transcript = True
            st.session_state["manual_input"] = "Goodbye. Thank you for coming."
            # Trigger the manual input immediately
            st.rerun()
else:
    if container3.button("Start Over", icon=":material/restart_alt:"):
        del st.session_state["text_chat_enabled"]
        st.session_state.messages = [st.session_state.messages[0]]
        st.rerun()

# Show the download button
if st.session_state.download_transcript:
    show_download(container3)
