import streamlit as st
from docx import Document
from io import BytesIO
from Session import push_session_log
from Utils import local_css
from Settings import settings
from mistletoe import markdown
from fpdf import FPDF, HTMLMixin
from fpdf.enums import AccessPermission

class MyFPDF(FPDF, HTMLMixin):
    pass


def create_transcript_pdf():
    md_lines = ["# Conversation Transcript", ""]
    debrief = False
    for msg in st.session_state.messages[1:]:
        speaker = (
            settings["user_name"]
            if msg["role"] == "user"
            else settings["assistant_name"]
        )
        if msg["role"] == "system":
            if not debrief:
                md_lines.append(f"---\n\n# Debrief\n\n")
            debrief = True
            speaker = "Instructor"
        if debrief and msg["role"] == "assistant":
            speaker = "Instructor"
        md_lines.append(f"**{speaker}:** {msg['content']}")
        md_lines.append("")  # blank line -> new paragraph

    md_text = "\n".join(md_lines)
    html = markdown(md_text)
    pdf = MyFPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_font("DejaVuSans", fname="assets/fonts/DejaVuSans.ttf")
    pdf.add_font("DejaVuSans", fname="assets/fonts/DejaVuSans-Bold.ttf", style="B")
    pdf.set_font("DejaVuSans", size=16)
    pdf.add_page()
    pdf.set_encryption(
        owner_password="drkim@32",
        permissions=AccessPermission.none()
    )
    pdf.write_html(html)
    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


def create_transcript_document():
    doc = Document()
    doc.add_heading("Conversation Transcript\n", level=1)

    for message in st.session_state.messages[1:]:
        if message["role"] == "user":
            p = doc.add_paragraph()
            p.add_run(settings["user_name"] + ": ").bold = True
            p.add_run(message["content"])
        else:
            p = doc.add_paragraph()
            p.add_run(settings["assistant_name"] + ": ").bold = True
            p.add_run(message["content"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


st.set_page_config(
    page_title="Download | " + settings["title"],
    page_icon=":material/download:",
    layout="wide",
    initial_sidebar_state="expanded",
)
st.title("Download | " + settings["title"])
# Inject CSS for custom styles
local_css("style.css")

st.markdown(settings["agreement"])

if st.checkbox("I agree to participate in the research."):
    push_session_log()

document = create_transcript_pdf()
st.download_button(
    label="Download Transcript",
    icon=":material/download:",
    data=document,
    file_name="Transcript.pdf",
    mime="application/pdf",
)
