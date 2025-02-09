import streamlit as st
import os
from openai import OpenAI
import anthropic
from io import BytesIO
import base64
from streamlit_mic_recorder import mic_recorder
from streamlit.runtime import get_instance
from streamlit.runtime.scriptrunner import get_script_run_ctx
import re
from docx import Document
import tomllib
import hmac
import warnings
import io
import uuid
import time
from datetime import datetime, timedelta
from Logger import get_logger
import codecs

warnings.filterwarnings("ignore", category=DeprecationWarning)


def get_uuid():
    timestamp = time.time()
    id = uuid.uuid4()
    id = f"{id}-{timestamp}"
    id = base64.urlsafe_b64encode(id.encode("utf-8")).decode("utf-8")
    return id[:8]


def get_session():
    runtime = get_instance()
    ctx = get_script_run_ctx()
    session_id = ctx.session_id
    session_info = runtime._instance.get_client(session_id)
    return session_id


@st.cache_resource
def get_active_users():
    return {}


def update_active_users(timeout=300):
    now = time.time()
    active_users = get_active_users()
    id = get_session()
    active_users[id] = now
    for user_id, last_active in list(active_users.items()):
        if now - last_active > timeout:
            del active_users[user_id]


def get_active_user_count():
    return len(get_active_users())


def elapsed(start):
    duration = time.time() - start
    duration_td = timedelta(seconds=duration)
    days = duration_td.days
    hours, remainder = divmod(duration_td.seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    dur_str = ""
    if days:
        dur_str = f"{days} days "
    if hours:
        dur_str += f"{hours} hours "
    if minutes:
        dur_str += f"{minutes} minutes "
    if seconds:
        dur_str += f"{seconds} seconds"
    return dur_str


@st.cache_resource
def get_openai_client():
    return OpenAI(api_key=st.secrets["OPENAI_API_KEY"])


@st.cache_resource
def get_anthropic_client():
    return anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])


def password_entered():
    """Checks whether a password entered by the user is correct."""
    if hmac.compare_digest(st.session_state["password"], st.secrets["password"]):
        st.session_state["password_correct"] = True
        del st.session_state["password"]  # Don't store the password.
        autoplay_audio(open("assets/unlock.mp3", "rb").read())
        log.info(f"Session Start: {get_session()}")
    else:
        st.session_state["password_correct"] = False


@st.cache_data
def load_settings():
    return tomllib.load(open("settings.toml", "rb"))


@st.cache_data
def local_css(file_name):
    with open(file_name) as f:
        st.html(f"<style>{f.read()}</style>")


@st.cache_data
def get_prompt():
    return codecs.open("prompt.txt", "r", "utf-8").read()


def speech_to_text(client, audio):
    try:
        id = audio["id"]
        log.debug(f"STT: {id}")
        audio_bio = io.BytesIO(audio["bytes"])
        audio_bio.name = "audio.wav"
        transcript = client.audio.transcriptions.create(
            model="whisper-1", response_format="text", file=audio_bio
        )
        st.session_state.processed_audio = id
        return transcript
    except Exception as e:
        log.exception("")


def text_to_speech(client, text):
    try:
        log.debug(f"TTS: {text}")
        response = client.audio.speech.create(
            model="tts-1",
            voice=st.session_state.settings["parameters"]["voice"],
            input=text,
        )
        autoplay_audio(response.content)
    except Exception as e:
        log.exception("")


def autoplay_audio(audio_data):
    b64 = base64.b64encode(audio_data).decode("utf-8")
    html_str = f"""
    <audio autoplay>
    <source src="data:audio/mp3;base64,{b64}" type="audio/mp3">
    </audio>
    """
    st.html(html_str)


# Send prompt to Anthropic and get response
def stream_response_anthropic(client, messages):
    try:
        log.debug(f"Sending text request to Anthropic: {messages[-1]['content']}")
        stream = client.messages.create(
            model=st.session_state.settings["parameters"]["model"],
            messages=messages[1:],
            max_tokens=1000,
            system=messages[0]["content"],
            temperature=st.session_state.settings["parameters"]["temperature"],
            stream=True,
        )
        for chunk in stream:
            if isinstance(
                chunk,
                anthropic.types.raw_content_block_delta_event.RawContentBlockDeltaEvent,
            ):
                yield chunk.delta.text
    except Exception as e:
        log.exception("")


# Send prompt to OpenAI and get response
def stream_response_openai(client, messages):
    try:
        log.debug(f"Sending text request to OpenAI: {messages[-1]['content']}")
        stream = client.chat.completions.create(
            model=st.session_state.settings["parameters"]["model"],
            messages=messages,
            temperature=st.session_state.settings["parameters"]["temperature"],
            stream=True,
        )
        for chunk in stream:
            if chunk.choices[0].delta.content is not None:
                yield chunk.choices[0].delta.content
    except Exception as e:
        log.exception("")


def show_download():
    document = create_transcript_document()
    col1, col2 = st.columns([1, 1])
    # Button to download the full conversation transcript
    with col1:
        st.download_button(
            label="📥 Download Transcript",
            data=document,
            file_name="Transcript.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


def create_transcript_document():
    doc = Document()
    doc.add_heading("Conversation Transcript\n", level=1)

    for message in st.session_state.messages[1:]:
        if message["role"] == "user":
            p = doc.add_paragraph()
            p.add_run(st.session_state.settings["user_name"] + ": ").bold = True
            p.add_run(message["content"])
        else:
            p = doc.add_paragraph()
            p.add_run(st.session_state.settings["assistant_name"] + ": ").bold = True
            p.add_run(message["content"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Session Initialization
def init_session():
    if "settings" not in st.session_state:
        st.session_state.settings = load_settings()
        defaults = {
            "show_intro": True,
            "chat_active": False,
            "messages": [{"role": "system", "content": get_prompt()}],
            "processed_audio": None,
            "manual_input": None,
            "end_session_button_clicked": False,
            "download_transcript": False,
            "start_time": time.time(),
        }
        for key, val in defaults.items():
            if key not in st.session_state:
                st.session_state[key] = val


def setup_sidebar():
    st.sidebar.header("Chat with " + st.session_state.settings["assistant_name"])
    # Button Login
    if "password_correct" in st.session_state and st.session_state.password_correct:
        st.session_state.chat_active = True
        st.session_state.show_intro = False
    else:
        st.sidebar.header("Access Code")
        with st.sidebar.container(border=True):
            with st.form("Credentials"):
                st.text_input("Access Code", type="password", key="password")
                st.form_submit_button("Start Chat", on_click=password_entered)
            if "password_correct" in st.session_state:
                st.error("😕 Invalid Code")

    if st.session_state.chat_active:
        # Update current user and display active users.
        update_active_users()
        if st.sidebar.button(f"🟢 Active Users: {get_active_user_count()}"):
            st.rerun()

        # Show Profile
        container1 = st.sidebar.container(border=True)
        with container1:
            for key, val in st.session_state.settings["sidebar"].items():
                if re.search(r"(jpg|png|webp)$", val):
                    st.image(val)
                else:
                    st.subheader(f"{key.replace("_", " ")}: {val}")


def show_messages():
    for message in st.session_state.messages[1:]:
        name = (
            st.session_state.settings["user_name"]
            if message["role"] == "user"
            else st.session_state.settings["assistant_name"]
        )
        avatar = (
            st.session_state.settings["user_avatar"]
            if message["role"] == "user"
            else st.session_state.settings["assistant_avatar"]
        )
        with st.chat_message(name, avatar=avatar):
            st.markdown(message["content"])


def handle_audio_input(client):
    with st.sidebar.container(border=True):
        audio = mic_recorder(
            start_prompt="🎙 Record",
            stop_prompt="📤 Stop",
            just_once=False,
            use_container_width=True,
            format="wav",
            key="recorder",
        )
    # Check if there is a new audio recording and it has not been processed yet
    if audio and audio["id"] != st.session_state.processed_audio:
        transcript = speech_to_text(client, audio)
        return transcript


def process_user_query(text_client, speech_client, user_query):
    # Display the user's query
    with st.chat_message(
        st.session_state.settings["user_name"],
        avatar=st.session_state.settings["user_avatar"],
    ):
        st.markdown(user_query)

    # Store the user's query into the history
    st.session_state.messages.append({"role": "user", "content": user_query.strip()})

    stream = (
        stream_response_openai(text_client, st.session_state.messages)
        if st.session_state.settings["parameters"]["model"].startswith("gpt")
        else stream_response_anthropic(text_client, st.session_state.messages)
    )
    response = "".join(stream).strip()
    if not st.session_state.end_session_button_clicked:
        text_to_speech(speech_client, response)

    with st.chat_message(
        st.session_state.settings["assistant_name"],
        avatar=st.session_state.settings["assistant_avatar"],
    ):
        st.markdown(response)

        st.session_state.messages.append(
            {"role": "assistant", "content": response}
        )


def main():
    # Inject CSS for custom styles
    local_css("style.css")
    init_session()
    text_client = (
        get_openai_client()
        if st.session_state.settings["parameters"]["model"].startswith("gpt")
        else get_anthropic_client()
    )
    speech_client = get_openai_client()
    st.title(st.session_state.settings["title"])
    setup_sidebar()

    # Display text before Start Chat is pressed
    if st.session_state.show_intro:
        with st.container(border=True):
            st.markdown(st.session_state.settings["intro"])

    # Check if chat is active
    if st.session_state.chat_active:
        show_messages()

        # Check if there's a manual input and process it
        if st.session_state.manual_input:
            user_query = st.session_state.manual_input
        else:
            if st.session_state.end_session_button_clicked:
                user_query = st.chat_input(
                    "Questions about your feedback? Ask them here."
                )
            else:
                user_query = st.chat_input(
                    "Click 'End Session' Button to Receive Feedback and Download Transcript."
                )
                transcript = handle_audio_input(speech_client)
                if transcript:
                    user_query = transcript

        if user_query:
            process_user_query(text_client, speech_client, user_query)
            if st.session_state.manual_input:
                st.session_state.manual_input = None
                st.rerun()

        # Handle end session
        if (
            not st.session_state.end_session_button_clicked
            and len(st.session_state.messages) > 1
        ):
            if st.button("End Session"):
                st.session_state.end_session_button_clicked = True
                log.info(
                    f"Session end: {elapsed(st.session_state.start_time)} {get_session()}"
                )
                st.session_state.download_transcript = True
                st.session_state["manual_input"] = "Goodbye. Thank you for coming."
                # Trigger the manual input immediately
                st.rerun()

        # Show the download button
        if st.session_state.download_transcript:
            show_download()

    st.sidebar.warning(st.session_state.settings["warning"])


# Configure page
st.set_page_config(
    page_title="AI SimBot",
    page_icon="💬",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Setup Log
if "logger" not in st.session_state:
    settings = load_settings()
    if "timezone" in settings:
        st.session_state.logger = get_logger(settings["timezone"])
    else:
        st.session_state.logger = get_logger()
log = st.session_state.logger

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        id = get_uuid()
        log.exception(f"Unhandled exception: {id}")
        st.error(
            f"{st.session_state.settings['error_message']}\n\n**Reference id: {id}**"
        )
